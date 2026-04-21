#!/usr/bin/env python3
"""Genera el documento de requerimientos de software para Hoteleria BCA."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(21.59)  # Letter
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Helper
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return h

def add_req(code, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{code}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.name = 'Arial'
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = 'Arial'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Arial'
    return table

# ==================== TITLE PAGE ====================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Especificacion de Requerimientos de Software")
run.bold = True
run.font.size = Pt(26)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sistema de Gestion Hotelera")
run.font.size = Pt(18)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x37, 0x56, 0xEB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Campamento BCA - Aramark Chile")
run.font.size = Pt(16)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

for _ in range(4):
    doc.add_paragraph()

info = [
    ("Version:", "1.1"),
    ("Fecha:", "Abril 2026"),
    ("Procedimiento base:", "PR-MEL-OPFC-001 v12 (31-03-2026)"),
    ("Preparado por:", "Equipo de Desarrollo - Daniel Hazan"),
    ("Cliente:", "Aramark Chile - Pablo Sierra"),
    ("Estado:", "Borrador para revision"),
]
for label, val in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Arial'
    r2 = p.add_run(val)
    r2.font.size = Pt(11)
    r2.font.name = 'Arial'

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
add_heading_styled("Tabla de Contenidos", 1)
toc_items = [
    "1. Introduccion",
    "   1.1 Proposito",
    "   1.2 Alcance",
    "   1.3 Definiciones y Acronimos",
    "2. Descripcion General",
    "   2.1 Perspectiva del Producto",
    "   2.2 Tipos de Usuario",
    "   2.3 Restricciones",
    "3. Requerimientos Funcionales",
    "   3.1 Gestion de Alojamiento",
    "   3.2 Gestion de Empresas y Dotacion",
    "   3.3 Facility Management",
    "   3.4 Gestion de Tarjetas Electronicas",
    "   3.5 Registros Operacionales",
    "   3.6 Conciliacion",
    "   3.7 Reporteria",
    "4. Requerimientos No Funcionales",
    "5. Matriz de Roles y Permisos",
    "6. Casos de Uso Principales",
    "7. Glosario",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.name = 'Arial'

doc.add_page_break()

# ==================== 1. INTRODUCCION ====================
add_heading_styled("1. Introduccion", 1)

add_heading_styled("1.1 Proposito", 2)
doc.add_paragraph(
    "El presente documento tiene como proposito definir los requerimientos funcionales y no funcionales "
    "para el desarrollo de la plataforma de Gestion Hotelera del Campamento BCA, operado por Aramark Chile. "
    "Este documento servira como base para el diseno, desarrollo e implementacion del sistema."
)

add_heading_styled("1.2 Alcance", 2)
doc.add_paragraph(
    "El sistema es una plataforma web unificada que reemplazara los 8 a 10 procesos manuales basados "
    "en planillas Excel que actualmente gestionan la operacion hotelera del campamento mas grande de Chile. "
    "El campamento alberga aproximadamente 3,700 personas de 162+ empresas con 55+ tipos de turno diferentes."
)
doc.add_paragraph(
    "El proyecto se alinea con el procedimiento oficial PR-MEL-OPFC-001 v12 (Asignacion de Habitaciones, "
    "31-03-2026) y soporta los campamentos 5400 (BCA), VMH, VVL y VCA."
)
doc.add_paragraph("El sistema cubrira los siguientes modulos:")
modules = [
    "Gestion de Alojamiento (check-in/out, asignacion de habitaciones, ocupacion)",
    "Gestion de Empresas y Dotacion (portal de autoservicio, curva semanal, cupos)",
    "Facility Management (ordenes de trabajo, limpieza, inspeccion)",
    "Gestion de Tarjetas Electronicas (llaves de acceso Keyplus)",
    "Registros Operacionales (censos, seguridad, quejas)",
    "Integracion y Conciliacion (OnTracking, SAMTECH, Keyplus, SharePoint)",
    "Reporteria y Analitica (incluye KPI-SharePoint de excepciones)",
]
for m in modules:
    doc.add_paragraph(m, style='List Bullet')

add_heading_styled("1.3 Definiciones y Acronimos", 2)
add_table(
    ["Termino", "Definicion"],
    [
        ["RUT", "Rol Unico Tributario - Identificador fiscal chileno"],
        ["OnTracking", "Sistema actual de asignacion de habitaciones y censos"],
        ["SAMTECH", "Sistema corporativo para registro de check-ins y censos"],
        ["Keyplus", "Sistema de cerraduras electronicas (origen chino)"],
        ["SharePoint", "Repositorio donde se publica el KPI de excepciones (dobles asignaciones)"],
        ["PR-MEL-OPFC-001", "Procedimiento oficial de asignacion de habitaciones v12 (31-03-2026)"],
        ["Curva", "Gestor que define y autoriza el ingreso de personal al campamento"],
        ["CO / ADC", "Contract Owner y Administrador de Contrato por parte del mandante"],
        ["Focal Point", "Contraparte en la empresa que coordina la curva y dotacion"],
        ["Cama P/V", "Cama Principal / Vacante en la asignacion de habitacion"],
        ["OPEX / CAPEX", "Clasificacion del centro de costo del contrato"],
        ["Near-miss", "Incidente de seguridad potencial (ej: asignacion de habitacion mixta)"],
        ["Doble Asignacion", "Excepcion donde una habitacion queda asignada a dos personas; se registra en KPI-SharePoint"],
        ["OT", "Orden de Trabajo"],
        ["BCA / 5400", "Campamento principal operado por Aramark; codigo interno 5400"],
        ["VMH / VVL / VCA", "Otros campamentos soportados por la plataforma"],
        ["Dotacion", "Lista semanal de personal que cada empresa envia"],
        ["Cupo", "Plaza de alojamiento asignada a una empresa"],
        ["Ineficiencia", "Cama desocupada o subutilizada en el campamento"],
        ["Regla 96h", "Anticipacion minima para solicitud de reserva (cierre Martes 18:00)"],
    ]
)

doc.add_page_break()

# ==================== 2. DESCRIPCION GENERAL ====================
add_heading_styled("2. Descripcion General", 1)

add_heading_styled("2.1 Perspectiva del Producto", 2)
doc.add_paragraph(
    "La plataforma sera un sistema web unificado que integrara los procesos actualmente dispersos "
    "en multiples planillas Excel, el sistema OnTracking (asignacion de habitaciones), el sistema "
    "SAMTECH (registro de check-ins / censos) y el sistema Keyplus (cerraduras electronicas). "
    "Reemplazara las aplicaciones locales desarrolladas por Javier Rivera, profesionalizandolas e "
    "integrandolas en una solucion estandarizada a nivel de compania, alineada con el procedimiento "
    "oficial PR-MEL-OPFC-001 v12."
)

add_heading_styled("2.2 Tipos de Usuario del Campamento", 2)
add_table(
    ["Tipo", "Descripcion", "Autorizacion Requerida"],
    [
        ["Permanente", "Contrato superior a 6 meses. Ocupa habitacion de forma continua, incluso con sistema de turno.", "Curva de la empresa"],
        ["Reemplazo", "Enviado por la empresa para ocupar el cupo de un titular en vacaciones o licencia.", "Curva de la empresa"],
        ["Visita", "Acude al campamento por uno o varios dias (ej: revision de sistema). Requiere doble autorizacion.", "Curva + Hoteleria"],
    ]
)

add_heading_styled("2.3 Restricciones", 2)
restrictions = [
    "El sistema Keyplus (cerraduras electronicas) no exporta el RUT en sus reportes, lo que impide la conciliacion automatica directa. Se debe utilizar habitacion + timestamp como alternativa.",
    "Los reportes de Keyplus estan limitados a 1,000 columnas por informe.",
    "La conectividad en el campamento puede ser limitada, requiriendo una solucion resiliente.",
    "Existen 55+ tipos de turno diferentes (7x7, 10x10, 14x14, 4x3, 8x6, etc.) que generan entre 600 y 700 ineficiencias diarias.",
    "Se manejan 162+ empresas con permisos y cupos distintos.",
    "Las solicitudes de reserva tienen una anticipacion minima de 96 horas; el cierre de curva semanal es el Martes a las 18:00.",
    "Las dobles asignaciones son excepciones que deben registrarse en el KPI-SharePoint segun el procedimiento PR-MEL-OPFC-001 v12.",
    "El sistema debe operar para multiples campamentos (5400 BCA, VMH, VVL, VCA) con configuracion independiente por sitio.",
    "El proceso debe ser lo suficientemente simple para que los recepcionistas lo usen sin errores.",
]
for r in restrictions:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ==================== 3. REQUERIMIENTOS FUNCIONALES ====================
add_heading_styled("3. Requerimientos Funcionales", 1)

# 3.1
add_heading_styled("3.1 Gestion de Alojamiento", 2)
reqs_31 = [
    ("RF-1.1", "El sistema debe implementar el check-in segun el flujo PR-MEL-OPFC-001 v12 con las siguientes etapas: (a) validacion del huesped y 96h de anticipacion, (b) asignacion de habitacion y cama P/V, (c) chequeo de condiciones (aseo, sabanas, casillero, infraestructura), (d) registro en sistemas (Planilla de Reservas, SAMTECH/SharePoint, enrolamiento Keyplus), (e) onboarding (Reglamento Interno, induccion, entregas)."),
    ("RF-1.2", "El sistema debe capturar los datos de check-in: RUT, nombre, sexo, contacto, cargo, tipo de usuario, empresa, tipo de contrato (Permanente/Spot/SD-PM/Outage/Proyecto/LSTS), centro de costo y clasificacion OPEX/CAPEX."),
    ("RF-1.3", "El sistema debe registrar habitacion, cama P/V, casillero asignado, fecha/hora de llegada y bajada, sistema de turno (4x3, 7x7, 8x6, 10x10, 14x14) y tipo de turno (Dia/Noche)."),
    ("RF-1.4", "El sistema debe soportar check-out manual y check-out automatico basado en eventos de OnTracking/Keyplus, y mantener una bandeja de pendientes de revision en terreno."),
    ("RF-1.5", "La revision en terreno durante el check-out debe incluir: devolucion de tarjeta Keyplus, devolucion de llave de casillero, inspeccion de habitacion, pertenencias, registro de danos y liberacion de casillero."),
    ("RF-1.6", "Al realizar un check-out, el sistema debe generar automaticamente las derivaciones correspondientes: OT de Limpieza, cambio de sabanas (segun ciclo del turno) y OT de Facility si se detectan danos."),
    ("RF-1.7", "El sistema debe asignar habitaciones automaticamente respetando las reglas de restriccion de genero establecidas por modulo."),
    ("RF-1.8", "El sistema debe gestionar el caso de excepcion 'Doble Asignacion', bloqueandolo por defecto y, si es autorizado, registrando la ocurrencia en el KPI-SharePoint."),
    ("RF-1.9", "El sistema debe mostrar un dashboard de ocupacion en tiempo real, desglosado por campamento, modulo, empresa y tipo de turno, incluyendo el KPI diario de dobles asignaciones."),
    ("RF-1.10", "El sistema debe detectar y visualizar las ineficiencias diarias (estimadas en 600-700 camas), identificando habitaciones con ocupacion inferior al 100%."),
    ("RF-1.11", "El sistema debe presentar un mapa visual del campamento mostrando el estado de ocupacion de cada modulo y habitacion con codigos de color."),
    ("RF-1.12", "El sistema debe gestionar los estados de habitacion: ocupada, disponible, reservada y bloqueada, permitiendo transiciones controladas."),
]
for code, text in reqs_31:
    add_req(code, text)

# 3.2
add_heading_styled("3.2 Gestion de Empresas y Dotacion", 2)
reqs_32 = [
    ("RF-2.1", "El sistema debe proveer un portal de autoservicio para las 162+ empresas donde puedan gestionar su dotacion, cupos y usuarios, con selector de campamento (5400/VMH/VVL/VCA)."),
    ("RF-2.2", "El portal debe mostrar la ficha de la empresa con: Razon Social, N° de Contrato, Gerencia General, Area, Contract Owner (CO), ADC, Focal Point Empresa y clasificacion OPEX/CAPEX."),
    ("RF-2.3", "El sistema debe permitir la carga semanal de dotacion mediante archivo Excel o via API, eliminando la intervencion manual del recepcionista."),
    ("RF-2.4", "El sistema debe aplicar la regla de 96 horas de anticipacion para solicitudes de reserva, con corte de curva semanal el Martes 18:00 y bloqueo de cargas fuera de plazo (salvo excepcion autorizada)."),
    ("RF-2.5", "Toda solicitud debe incluir: ID auto-generado, campamento destino, sistema de turno, tipo de turno, tipo de contrato, centro de costo y clasificacion OPEX/CAPEX."),
    ("RF-2.6", "El sistema debe validar automaticamente los datos de dotacion cargados, verificando consistencia de genero, cupos disponibles y formatos de RUT."),
    ("RF-2.7", "El sistema debe gestionar los cupos por empresa con un mecanismo de recuperacion de cupos subutilizados, permitiendo reasignarlos a otras empresas."),
    ("RF-2.8", "El sistema debe soportar formularios QR diferenciados por empresa (hasta 155 formularios) para agilizar el proceso de check-in."),
    ("RF-2.9", "El portal debe mostrar la curva de poblamiento proyectada, la dotacion semanal, el historial de solicitudes y los usuarios habilitados."),
]
for code, text in reqs_32:
    add_req(code, text)

# 3.3
add_heading_styled("3.3 Facility Management", 2)
reqs_33 = [
    ("RF-3.1", "El sistema debe permitir la creacion de ordenes de trabajo categorizadas por tipo: limpieza, reparacion, cambio de colchon, cambio de cortina, reposicion de insumos y casilleros."),
    ("RF-3.2", "El sistema debe recibir automaticamente OTs generadas desde el check-out (limpieza y sabanas) y desde la revision de terreno (casillero liberado, danos)."),
    ("RF-3.3", "El sistema debe asignar ordenes de trabajo a personal especifico y generar alertas automaticas cuando el tiempo de respuesta exceda los umbrales definidos (ej: 48 horas)."),
    ("RF-3.4", "El sistema debe monitorear la productividad del auxiliar de aseo, registrando camas hechas, reposiciones de insumos y limpiezas realizadas."),
    ("RF-3.5", "El sistema debe controlar los cambios de sabanas segun el sistema de turno de cada usuario, informando al area de Facility para la planificacion del servicio."),
    ("RF-3.6", "El sistema debe permitir la inspeccion de habitaciones (Camp Check) evaluando disponibilidad, limpieza, banos, infraestructura e insumos, con historial consultable."),
]
for code, text in reqs_33:
    add_req(code, text)

# 3.4
add_heading_styled("3.4 Gestion de Tarjetas Electronicas (Keyplus)", 2)
reqs_34 = [
    ("RF-4.1", "El sistema debe permitir el enrolamiento de tarjetas Keyplus durante el check-in, asociandolas al huesped, habitacion y casillero."),
    ("RF-4.2", "El sistema debe permitir la asignacion y devolucion de tarjetas maestras electronicas, registrando responsable, fecha y modulo de acceso."),
    ("RF-4.3", "El sistema debe soportar la importacion masiva de tarjetas desde archivos Excel."),
    ("RF-4.4", "El sistema debe mantener una trazabilidad historica completa de cada tarjeta: asignacion, devolucion, perdida y reposicion."),
    ("RF-4.5", "El sistema debe documentar las perdidas de tarjetas y controlar el acceso por modulo, previniendo que usuarios no autorizados identifiquen el modulo al que pueden acceder."),
    ("RF-4.6", "El sistema debe soportar la devolucion de tarjeta en la etapa de revision de terreno del check-out, confirmando la liberacion antes de cerrar el cupo."),
]
for code, text in reqs_34:
    add_req(code, text)

# 3.5
add_heading_styled("3.5 Registros Operacionales", 2)
reqs_35 = [
    ("RF-5.1", "El sistema debe registrar el censo diario de ocupacion desglosado por modulo, empresa y tipo de turno."),
    ("RF-5.2", "El sistema debe registrar eventos de seguridad ocurridos en el campamento."),
    ("RF-5.3", "El sistema debe registrar quejas y levantamientos de manos de los usuarios y empresas."),
    ("RF-5.4", "El sistema debe permitir la exportacion de registros en formato CSV y Excel con filtros de fecha y modulo."),
]
for code, text in reqs_35:
    add_req(code, text)

# 3.6
add_heading_styled("3.6 Integracion y Conciliacion", 2)
reqs_36 = [
    ("RF-6.1", "El sistema debe integrarse con OnTracking (asignaciones de habitaciones y censos), SAMTECH (registro de check-ins) y Keyplus (registros de apertura de cerraduras)."),
    ("RF-6.2", "El sistema debe realizar la comparacion automatica entre los tres sistemas, utilizando RUT cuando este disponible (OnTracking/SAMTECH) y habitacion + timestamp como clave alternativa para Keyplus."),
    ("RF-6.3", "El sistema debe visualizar claramente las discrepancias encontradas: registros en OnTracking sin apertura en Keyplus, aperturas en Keyplus sin registro en OnTracking/SAMTECH, y check-ins en SAMTECH sin reserva aprobada."),
    ("RF-6.4", "El sistema debe publicar el KPI de dobles asignaciones al SharePoint corporativo segun el procedimiento PR-MEL-OPFC-001 v12."),
    ("RF-6.5", "El sistema debe generar un reporte diario de conciliacion exportable."),
]
for code, text in reqs_36:
    add_req(code, text)

# 3.7
add_heading_styled("3.7 Reporteria", 2)
reqs_37 = [
    ("RF-7.1", "El sistema debe generar un reporte de censo diario consolidado."),
    ("RF-7.2", "El sistema debe generar reportes de ocupacion desglosados por modulo y empresa."),
    ("RF-7.3", "El sistema debe generar reportes de ineficiencias identificando camas vacias y subutilizadas."),
    ("RF-7.4", "El sistema debe generar reportes de dotacion por gerencia y proyecto."),
    ("RF-7.5", "El sistema debe proveer un dashboard ejecutivo con visualizaciones graficas y capacidad de exportacion a Excel/PDF."),
]
for code, text in reqs_37:
    add_req(code, text)

doc.add_page_break()

# ==================== 4. REQUERIMIENTOS NO FUNCIONALES ====================
add_heading_styled("4. Requerimientos No Funcionales", 1)

rnfs = [
    ("RNF-1 Rendimiento", "Las operaciones frecuentes (check-in, consultas, busquedas) deben responder en menos de 2 segundos. El dashboard debe actualizar datos en menos de 5 segundos."),
    ("RNF-2 Disponibilidad", "El sistema debe mantener una disponibilidad minima de 99.5% (uptime), considerando que la operacion es 24/7."),
    ("RNF-3 Seguridad", "El sistema debe implementar autenticacion segura, gestion de roles y permisos por modulo, y cifrado de datos sensibles (RUT, datos personales)."),
    ("RNF-4 Usabilidad", "La interfaz debe ser simple e intuitiva, especialmente para los recepcionistas que realizan las operaciones mas frecuentes. Debe ser responsive para uso en tablets."),
    ("RNF-5 Localizacion", "El sistema debe estar en espanol (Chile), soportar formato de RUT chileno, zona horaria America/Santiago y formato de fechas DD/MM/AAAA."),
    ("RNF-6 Integracion", "El sistema debe exponer una API REST para integracion con OnTracking y soportar importacion de archivos del sistema Keyplus."),
    ("RNF-7 Escalabilidad", "El sistema debe soportar 3,700+ usuarios registrados con al menos 200 usuarios concurrentes del sistema."),
    ("RNF-8 Auditoria", "El sistema debe mantener un log completo de todas las operaciones realizadas, incluyendo usuario, fecha/hora, accion y datos modificados."),
]
for title, desc in rnfs:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{title}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Arial'
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)
    r2.font.name = 'Arial'

doc.add_page_break()

# ==================== 5. MATRIZ DE ROLES ====================
add_heading_styled("5. Matriz de Roles y Permisos", 1)
doc.add_paragraph(
    "La siguiente tabla define los permisos de cada rol en los modulos del sistema. "
    "L = Lectura, E = Escritura, A = Administracion, - = Sin acceso."
)

roles_headers = ["Modulo", "Admin", "Recepcionista", "Facility Mgr", "Sup. Aseo", "Admin Empresa", "Gerencia"]
roles_data = [
    ["Alojamiento", "A", "E", "L", "-", "L", "L"],
    ["Empresas/Dotacion", "A", "E", "L", "-", "E", "L"],
    ["Facility / OT", "A", "L", "A", "E", "-", "L"],
    ["Camp Check", "A", "E", "E", "E", "-", "L"],
    ["Tarjetas", "A", "E", "-", "-", "-", "L"],
    ["Registros", "A", "E", "E", "E", "-", "L"],
    ["Conciliacion", "A", "E", "L", "-", "-", "L"],
    ["Reportes", "A", "L", "L", "L", "L", "A"],
    ["Usuarios", "A", "-", "-", "-", "-", "-"],
]
add_table(roles_headers, roles_data)

doc.add_page_break()

# ==================== 6. CASOS DE USO ====================
add_heading_styled("6. Casos de Uso Principales", 1)

use_cases = [
    ("CU-01: Check-in segun PR-MEL-OPFC-001 v12",
     "Actor: Recepcionista. Flujo de 6 pasos: (1) validacion del huesped (RUT, autorizacion, regla 96h); (2) asignacion de "
     "habitacion, cama P/V y casillero; (3) chequeo de condiciones (aseo, sabanas, casillero, infraestructura); "
     "(4) registro en Planilla de Reservas, SAMTECH/SharePoint y enrolamiento de tarjeta Keyplus; (5) entrega del "
     "Reglamento Interno, induccion y kit; (6) actualizacion de ocupacion en tiempo real."),
    ("CU-02: Check-in de usuario de visita",
     "Actor: Recepcionista. Similar a CU-01 pero requiere verificar que existe autorizacion de la Curva de la empresa "
     "Y autorizacion del area de Hoteleria. Sin ambas autorizaciones el check-in es rechazado."),
    ("CU-03: Check-out automatico con revision de terreno",
     "Actor: Recepcionista / Sistema. El sistema detecta el check-out por eventos OnTracking/Keyplus o registro manual, "
     "agrega la habitacion a la bandeja de 'Pendientes Revision Terreno'. Al completar la revision (tarjeta, llave casillero, "
     "inspeccion, pertenencias, danos, casillero liberado), el sistema genera automaticamente OTs de limpieza, sabanas y "
     "Facility si corresponde."),
    ("CU-04: Excepcion Doble Asignacion",
     "Actor: Supervisor. Cuando se detecta el intento de asignar una habitacion ya ocupada, el sistema bloquea la operacion "
     "y solicita autorizacion. Si es autorizada, se registra el incidente en el KPI-SharePoint y se notifica para reasignacion inmediata."),
    ("CU-05: Solicitud semanal de dotacion con regla 96h",
     "Actor: Focal Point Empresa. La empresa accede al portal seleccionando campamento (5400/VMH/VVL/VCA), completa la curva "
     "con sistema de turno, tipo de contrato, centro de costo y clasificacion OPEX/CAPEX. El sistema valida la regla de 96h "
     "(corte Martes 18:00) y automaticamente los datos, notificando discrepancias."),
    ("CU-06: Recuperacion de cupos subutilizados",
     "Actor: Recepcionista/Admin. El sistema identifica cupos que no han sido utilizados en X dias. El recepcionista "
     "puede recuperar estos cupos y reasignarlos a empresas que necesitan aumentar su dotacion de permanentes."),
    ("CU-07: Creacion y seguimiento de orden de trabajo",
     "Actor: Cualquier usuario/empresa. Se crea una OT especificando tipo (limpieza, reparacion, etc.), ubicacion y "
     "descripcion. El Facility Manager la asigna a un auxiliar. El sistema genera alertas si la OT excede el tiempo limite."),
    ("CU-08: Conciliacion diaria OnTracking / SAMTECH / Keyplus",
     "Actor: Recepcionista/Admin. Se importan los archivos de los tres sistemas. El sistema cruza los datos por RUT y por "
     "habitacion + timestamp, mostrando discrepancias para revision manual y publicando el KPI al SharePoint."),
    ("CU-09: Inspeccion de habitacion (Camp Check)",
     "Actor: Recepcionista. Se selecciona el modulo y habitacion, se evalua disponibilidad, limpieza, banos, infraestructura "
     "e insumos con calificacion Bueno/Regular/Malo. Se registran observaciones. El historial queda consultable."),
    ("CU-10: Generacion de reporte de censo diario",
     "Actor: Gerencia. Se selecciona la fecha y filtros deseados. El sistema genera el reporte consolidado con "
     "ocupacion por campamento, modulo, empresa y turno. Se puede exportar a Excel o PDF."),
]
for title, desc in use_cases:
    p = doc.add_paragraph()
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Arial'
    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.space_after = Pt(12)
    for r in p2.runs:
        r.font.size = Pt(11)
        r.font.name = 'Arial'

doc.add_page_break()

# ==================== 7. GLOSARIO ====================
add_heading_styled("7. Glosario", 1)

glossary = [
    ("RUT", "Rol Unico Tributario. Identificador fiscal unico de cada persona en Chile. Formato: XX.XXX.XXX-X."),
    ("OnTracking", "Sistema web utilizado actualmente para registrar asignaciones de habitaciones y generar censos de ocupacion. Requiere entrada manual de datos."),
    ("SAMTECH", "Sistema corporativo utilizado para registrar check-ins, apoyando el cruce de datos con OnTracking y Keyplus."),
    ("Keyplus", "Sistema de cerraduras electronicas utilizado en el campamento. No exporta RUT en sus reportes y tiene un limite de 1,000 columnas por informe."),
    ("PR-MEL-OPFC-001", "Procedimiento oficial de asignacion de habitaciones, version 12 vigente desde el 31-03-2026."),
    ("Campamentos", "El sistema soporta multiples campamentos: 5400 (BCA, principal), VMH, VVL y VCA."),
    ("Curva", "Gestor designado por cada empresa que define y autoriza el ingreso de personal al campamento. Controla los cupos y sistemas de turno."),
    ("CO / ADC", "Contract Owner y Administrador de Contrato designados por el mandante para cada contrato."),
    ("Focal Point", "Contraparte en la empresa contratista que centraliza la carga de curva y dotacion."),
    ("Cama P/V", "Cama Principal o Vacante. Define el rol de la cama en la asignacion dentro de la habitacion."),
    ("OPEX / CAPEX", "Clasificacion contable del centro de costo asociado al contrato (operacional vs inversion)."),
    ("Tipos de Contrato", "Permanente, Spot, SD/PM, Outage, Proyecto y LSTS."),
    ("Sistemas de Turno", "4x3, 7x7, 8x6, 10x10, 14x14 entre otros; cada uno con su patron dia/noche."),
    ("Regla 96h", "Anticipacion minima de 96 horas para solicitudes de reserva; corte de curva semanal los Martes a las 18:00."),
    ("Doble Asignacion", "Excepcion al procedimiento donde una misma habitacion queda asignada a dos huespedes; debe registrarse en el KPI-SharePoint."),
    ("Near-miss", "Incidente de seguridad potencial que no llego a materializarse. En hoteleria, se refiere tipicamente a la asignacion incorrecta de una habitacion que podria resultar en cruce de generos."),
    ("Dotacion", "Lista semanal que cada empresa envia con el personal que requiere alojamiento, incluyendo nombres, RUT, turnos y tipo de usuario."),
    ("Cupo", "Plaza de alojamiento asignada a una empresa. Puede ser para usuario permanente, reemplazo o visita."),
    ("Ineficiencia", "Cama desocupada o subutilizada en el campamento. Se estiman entre 600 y 700 ineficiencias diarias debido a la mezcla de 55+ tipos de turno."),
    ("Levantamiento de manos", "Termino utilizado para referirse a quejas formales registradas por usuarios o empresas sobre el servicio."),
    ("OT", "Orden de Trabajo. Solicitud de servicio al area de Facility (limpieza, reparacion, cambio de colchon, casillero, etc.)."),
    ("Modulo", "Division fisica del campamento que agrupa un conjunto de habitaciones. Cada modulo tiene administracion y control de acceso independiente."),
]
add_table(["Termino", "Definicion"], glossary)

# Save
output_path = "/Users/sheilabriceno/Downloads/Hotelería/docs/requerimientos-hoteleria.docx"
doc.save(output_path)
print(f"Documento generado: {output_path}")
